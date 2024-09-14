import os
import subprocess
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
import re
import fitz  # PyMuPDF
from docx import Document
import json
import nltk
from nltk.corpus import stopwords
from sklearn.linear_model import LogisticRegression
from sklearn.pipeline import Pipeline
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from transformers import pipeline
import spacy
from langdetect import detect
import xlrd  # Excel
import win32com.client  # Outlook
from PIL import Image
import pytesseract
import xml.etree.ElementTree as ET  # Visio
import fitz  # PyMuPDF
from PyPDF2 import PdfReader
from docx import Document
import xlrd
import win32com.client
import pytesseract
from PIL import Image
import xml.etree.ElementTree as ET
from sklearn.model_selection import GridSearchCV



model_weight= 1
keyword_weight= 0


param_grid = {
    'tfidf__max_df': [0.5, 0.75, 1.0],
    'tfidf__ngram_range': [(1, 1), (1, 2)],
    'clf__C': [1.0, 10.0, 100.0]
}

def parse_metadata(metadata):
    try:
        # Hier kannst du spezifische Logik für die Verarbeitung der Metadaten hinzufügen.
        # Falls die Metadaten bereits in einem lesbaren Format vorliegen, kannst du sie direkt zurückgeben.
        return metadata
    except Exception as e:
        print(f"Fehler beim Parsen der Metadaten: {e}")
        return {}


def grid_search_train(X_train, y_train, metrics_file):
    pipeline = Pipeline([
        ('tfidf', TfidfVectorizer(stop_words=german_stopwords)),
        ('clf', LogisticRegression(max_iter=200))
    ])

    # GridSearchCV initialisieren
    grid_search = GridSearchCV(pipeline, param_grid, scoring='accuracy', cv=5, n_jobs=-1, verbose=2)

    # Training mit Grid Search
    grid_search.fit(X_train, y_train)

    # Beste Parameter und bester Score
    best_params = grid_search.best_params_
    best_score = grid_search.best_score_

    # Schreiben der besten Parameter und des besten Scores in die metrics-Datei
    with open(metrics_file, 'a') as f:
        f.write(f"Beste Hyperparameter: {best_params}\n")
        f.write(f"Beste Accuracy: {best_score}\n")

    return grid_search.best_estimator_


# Laden der deutschen Stopwörter
nltk.download('stopwords')
german_stopwords = stopwords.words('german')

# Konfiguration aus der Datei laden
with open('config.json', 'r') as config_file:
    config = json.load(config_file)

# Verzeichnisse und Dateien aus der Konfiguration
documents_dir = config['documents_dir']
output_dir = config['output_dir']
output_file = os.path.join(output_dir, config['output_file'])
output_image = os.path.join(output_dir, config['output_image'])
training_directories = config['training_directories']
TLP_COLORS = config['tlp_colors']

# Prüfen, ob keywords.json existiert, falls nicht, ausführen
if not os.path.exists('keywords.json'):
    subprocess.run(['python', 'load_keywords.py'], check=True)

# Keywords aus keywords.json laden
with open('keywords.json', 'r', encoding='utf-8') as keyword_file:
    keyword_map = json.load(keyword_file)

# Spacy Modell für NER laden
nlp = spacy.load("de_core_news_sm")

# Vortrainiertes BERT-Modell für die Textklassifikation laden
classifier = pipeline('text-classification', model='oliverguhr/german-sentiment-bert', top_k=None)

# Funktion zum Extrahieren von Text aus PDF-Dateien
def extract_text_from_pdf(file_path):
    text = ""
    try:
        pdf_reader = PdfReader(file_path)
        for page in pdf_reader.pages:
            text += page.extract_text() if page.extract_text() else ''
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus DOCX-Dateien
def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus TXT-Dateien
def extract_text_from_txt(file_path):
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus Excel-Dateien
def extract_text_from_xlsx(file_path):
    text = ""
    try:
        workbook = xlrd.open_workbook(file_path)
        sheet = workbook.sheet_by_index(0)
        text = "\n".join(["\t".join(map(str, sheet.row_values(row))) for row in range(sheet.nrows)])
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus MSG-Dateien
def extract_text_from_msg(file_path):
    text = ""
    try:
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        msg = outlook.OpenSharedItem(file_path)
        text = msg.Body
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus Bild-Dateien
def extract_text_from_jpg(file_path):
    text = ""
    try:
        text = pytesseract.image_to_string(Image.open(file_path))
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Text aus Visio-Dateien
def extract_text_from_vsdx(file_path):
    text = ""
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        text = " ".join([elem.text for elem in root.iter() if elem.text])
    except Exception as e:
        print(f"Fehler beim Lesen von {file_path}: {e}")
    return text.lower()

# Funktion zum Extrahieren von Metadaten aus PDF- und DOCX-Dateien
def extract_metadata(file_path):
    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            metadata = doc.metadata
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            metadata = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "subject": doc.core_properties.subject
            }
        else:
            metadata = {}
        return metadata
    except Exception as e:
        return {}

# Funktion zum Extrahieren von Text und Metadaten aus allen Dateien in einem Verzeichnis
def extract_documents_text(directory):
    documents = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.startswith('~$'):
                continue
            path = os.path.join(root, file)
            cleaned_path = clean_path(path)
            metadata = extract_metadata(path)
            text = ""
            if file.endswith('.pdf'):
                text = extract_text_from_pdf(path)
            elif file.endswith('.docx'):
                text = extract_text_from_docx(path)
            elif file.endswith('.txt'):
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        text = f.read().lower()
                except Exception as e:
                    text = ""
            elif file.endswith('.xlsx'):
                text = extract_text_from_xlsx(path)
            elif file.endswith('.msg'):
                text = extract_text_from_msg(path)
            elif file.endswith('.jpg'):
                text = extract_text_from_jpg(path)
            elif file.endswith('.vsdx'):
                text = extract_text_from_vsdx(path)
            
            if text:
                documents.append((cleaned_path, text, metadata))
            else:
                print(f"Could not extract text from: {path}")
    return documents


# Pfad bereinigen (Normalisieren und Kleinbuchstaben)
def clean_path(path):
    return os.path.normpath(path).replace("\\", "/").lower()

# Funktion zum Extrahieren von Trainingsdaten aus den angegebenen Verzeichnissen
def extract_training_data(training_directories):
    training_documents = []
    for label, directory in training_directories.items():
        documents = extract_documents_text(directory)
        for _, text, metadata in documents:
            training_documents.append((directory, text, label, metadata))
    return training_documents

# Funktion zum Extrahieren von Features (TF-IDF) aus den Dokumenten
def extract_features(documents):
    texts = [text for _, text, _ in documents]
    metadata_texts = [" ".join([str(value) for value in metadata.values()]) for _, _, metadata in documents]
    combined_texts = [f"{text} {metadata}" for text, metadata in zip(texts, metadata_texts)]
    vectorizer = TfidfVectorizer(stop_words=german_stopwords)
    X = vectorizer.fit_transform(combined_texts)
    return X, vectorizer

# Funktion zur Klassifizierung auf Basis von Keywords
def keyword_classification(text, keyword_map, keyword_weight):
    scores = {label: 0 for label in keyword_map.keys()}
    for label, keywords in keyword_map.items():
        for keyword in keywords:
            if keyword.lower() in text:
                scores[label] += 1
    
    # Normalisieren der Scores
    max_score = max(scores.values())
    if max_score == 0:
        return 'UNKNOWN', 0
    for label in scores:
        scores[label] = (scores[label] / max_score) * keyword_weight
    
    best_label = max(scores, key=scores.get)
    return best_label, scores[best_label]

# Funktion zum Extrahieren von Entitäten aus dem Text mittels Spacy
def extract_entities(text):
    doc = nlp(text)
    entities = [ent.text for ent in doc.ents]
    return entities

# Funktion zur Sentimentanalyse mittels vortrainiertem BERT-Modell
def analyze_sentiment(text):
    results = classifier(text)
    sentiment_scores = {result['label']: result['score'] for result in results[0]}
    return sentiment_scores

# Erweiterte Klassifizierung durch Kombination von Keyword- und Modell-basierter Klassifizierung


#def enhanced_classification(text, metadata, keyword_map, model, keyword_weight, model_weight):
#    # Dynamischer Schwellenwert, initialisiert auf einen Basiswert
#    base_threshold = 0.5
#
#    # Vorbereitung des kombinierten Textes
#    combined_text = f"{text} {' '.join([str(value) for value in metadata.values()])}"
#
#    # Keyword-basierte Klassifizierung
#    keyword_label, keyword_score = keyword_classification(combined_text, keyword_map, keyword_weight)
#    keyword_score = max(keyword_score, 1e-5)  # Vermeidung von log(0) und Division durch Null
#
#    # Modell-basierte Klassifizierung
#    model_label = model.predict([combined_text])[0]
#    model_score = model_weight  # Wir gehen davon aus, dass dieser Score proportional zur Gewichtung ist
#
#    # Sicherstellen, dass keine Division durch Null stattfindet
#    if np.isnan(keyword_score) or np.isnan(model_score):
#        print("Warning: NaN detected in scores")
#        return 'UNKNOWN'
#
#    # Normalisierung der Scores auf einen Bereich von 0 bis 1
#    keyword_score = (keyword_score - np.min(keyword_score)) / (np.max(keyword_score) - np.min(keyword_score) + 1e-5)
#    model_score = (model_score - np.min(model_score)) / (np.max(model_score) - np.min(model_score) + 1e-5)
#
#    # Sicherstellen, dass die Normalisierung keine NaN-Werte erzeugt
#    if np.isnan(keyword_score) or np.isnan(model_score):
#        print("Warning: NaN detected after normalization")
#        return 'UNKNOWN'
#
#    # Nicht-lineare Kombination mit einer Sigmoid-Funktion
#    final_score = (keyword_weight * keyword_score + model_weight * model_score) / (keyword_weight + model_weight)
#    final_score = 1 / (1 + np.exp(-final_score))  # Sigmoid-Funktion zur Glättung
#
#    # Sicherstellen, dass der finale Score gültig ist
#    if np.isnan(final_score):
#        print("Warning: NaN detected in final score")
#        return 'UNKNOWN'
#
#    # Optional: Anpassung des Schwellenwerts basierend auf den Scores
#    adaptive_threshold = base_threshold * np.mean([keyword_score, model_score])
#
#    # Debugging-Ausgabe
#    print(f"Keyword Score: {keyword_score}, Model Score: {model_score}, Final Score: {final_score}, Threshold: {adaptive_threshold}")
#
#    # Entscheidung basierend auf dem finalen Score und dem adaptiven Schwellenwert
#    if final_score >= adaptive_threshold:
#        return model_label if model_score >= keyword_score else keyword_label
#    else:
#        return 'UNKNOWN'


def enhanced_classification(text, metadata, keyword_map, model, keyword_weight, model_weight):
    # Text und Metadaten kombinieren
    combined_text = f"{text} {' '.join([str(value) for value in metadata.values()])}"
    
    # Keyword-basierte Klassifizierung
    keyword_label, keyword_score = keyword_classification(combined_text, keyword_map, keyword_weight)
    
    # Modell-basierte Klassifizierung
    model_label = model.predict([combined_text])[0]
    model_score = model_weight  # Placeholder, hier können Modellwahrscheinlichkeiten oder andere Metriken verwendet werden

    # Ergebnisse kombinieren
    if keyword_score > model_score:
        return keyword_label
    else:
        return model_label

# Dokumente mit erweiterten Techniken klassifizieren
def classify_documents_with_enhanced_techniques(documents, keyword_map, model, keyword_weight, model_weight):
    classified_documents = []
    unclassified_documents = []
    for path, text, metadata in documents:
        label = enhanced_classification(text, metadata, keyword_map, model, keyword_weight, model_weight)
        if label == 'UNKNOWN':
            unclassified_documents.append((path, text, metadata))
        else:
            classified_documents.append((path, text, metadata, label))
    return classified_documents, unclassified_documents

# Sensitivitätsmodell trainieren
# Sensitivitätsmodell trainieren
# Trainingsfunktion anpassen
def train_sensitivity_model(training_documents, metrics_file):
    df = pd.DataFrame(training_documents, columns=['path', 'text', 'label', 'metadata'])
    
    # Metadaten korrekt parsen
    df['metadata'] = df['metadata'].apply(parse_metadata)

    valid_labels = set(TLP_COLORS.keys()) - {'UNKNOWN'}
    df = df[df['label'].isin(valid_labels)]

    X = df.apply(lambda row: f"{row['text']} {' '.join([str(value) for value in row['metadata'].values()])}", axis=1)
    y = df['label']

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=1, random_state=42)

    # Beste Modelltraining durch Grid Search
    best_model = grid_search_train(X_train, y_train, metrics_file)

    y_pred = best_model.predict(X_test)
    accuracy = accuracy_score(y_test, y_pred)
    print(f"Model accuracy mit besten Parametern: {accuracy}")

    # Schreiben der finalen Accuracy in die metrics-Datei
    with open(metrics_file, 'a') as f:
        f.write(f"Finale Model Accuracy: {accuracy}\n")

    return best_model


# Dokumente klassifizieren
def classify_documents(model, documents):
    combined_texts = [f"{text} {' '.join([str(value) for value in metadata.values()])}" for _, text, metadata in documents]
    labels = model.predict(combined_texts)
    return labels

# Cluster plotten und speichern
def plot_clusters(X, labels, output_file=output_image):
    pca = PCA(n_components=2)
    scatter_plot_points = pca.fit_transform(X.toarray())
    x_axis = scatter_plot_points[:, 0]
    y_axis = scatter_plot_points[:, 1]
    
    fig, ax = plt.subplots(figsize=(20, 10))
    unique_labels = set(labels)
    
    for label in unique_labels:
        indices = [i for i, l in enumerate(labels) if l == label]
        ax.scatter(x_axis[indices], y_axis[indices], label=label, color=TLP_COLORS.get(label, 'black'))
    
    ax.legend()
    plt.savefig(output_file)

# Klassifizierungsergebnisse speichern

def save_classification_results(documents, labels, output_file):
    color_to_label = {"0": 0, "1": 1, "2": 2, "3": 3}
    numeric_labels = [color_to_label.get(label, label) for label in labels]
    data = [(path, label) for (path, _, _), label in zip(documents, numeric_labels)]
    df = pd.DataFrame(data, columns=['Document Path', 'Sensitivity Label'])
    df.to_csv(output_file, index=False)

    data = [(path, label) for (path, _, _), label in zip(documents, labels)]
    df = pd.DataFrame(data, columns=['Document Path', 'Sensitivity Label'])
    df.to_csv(output_file, index=False)

# Hauptfunktion zur Durchführung des gesamten Klassifizierungsprozesses
def main_full(keyword_weight, model_weight):
    documents = extract_documents_text(documents_dir)
    if not documents:
        return
    
    training_documents = extract_training_data(training_directories)
    
    metrics_file = 'metrics.txt'  # Definiere den Pfad zur metrics-Datei
    
    sensitivity_model = train_sensitivity_model(training_documents, metrics_file)
    
    if sensitivity_model is None:
        return

    classified_documents, unclassified_documents = classify_documents_with_enhanced_techniques(documents, keyword_map, sensitivity_model, keyword_weight, model_weight)
    
    if unclassified_documents:
        X_unclassified, vectorizer = extract_features(unclassified_documents)
        kmeans = KMeans(n_clusters=4, random_state=42).fit(X_unclassified)
        unsupervised_labels = kmeans.labels_

        for (path, text, metadata), label in zip(unclassified_documents, unsupervised_labels):
            classified_documents.append((path, text, metadata, f'CLUSTER_{label}'))

    labels = [label for _, _, _, label in classified_documents]
    save_classification_results(documents, labels, output_file)
    X, vectorizer = extract_features(documents)
    plot_clusters(X, labels)

# Hauptprogramm ausführen
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
main_full(keyword_weight, model_weight)
