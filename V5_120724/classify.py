import os
import subprocess
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
import re
import fitz  # PyMuPDF
from docx import Document
import json
import nltk
from nltk.corpus import stopwords
from sklearn.linear_model import LogisticRegression
from sklearn.pipeline import Pipeline
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from transformers import pipeline
import spacy
from langdetect import detect

nltk.download('stopwords')
german_stopwords = stopwords.words('german')

# Load configuration from file
with open('config.json', 'r') as config_file:
    config = json.load(config_file)

documents_dir = config['documents_dir']
output_dir = config['output_dir']
output_file = os.path.join(output_dir, config['output_file'])
output_image = os.path.join(output_dir, config['output_image'])
training_directories = config['training_directories']
TLP_COLORS = config['tlp_colors']

# Check if keywords.json exists
if not os.path.exists('keywords.json'):
    subprocess.run(['python', 'load_keywords.py'], check=True)

# Load keywords from keywords.json
with open('keywords.json', 'r', encoding='utf-8') as keyword_file:
    keyword_map = json.load(keyword_file)

# Lade Spacy Modell für NER
nlp = spacy.load("de_core_news_sm")

# Laden des vortrainierten BERT-Modells für die Textklassifikation
classifier = pipeline('text-classification', model='oliverguhr/german-sentiment-bert', top_k=None)

def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text.lower()
    except Exception as e:
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.lower()
    except Exception as e:
        return ""

def extract_metadata(file_path):
    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            metadata = doc.metadata
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            metadata = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "subject": doc.core_properties.subject
            }
        else:
            metadata = {}
        return metadata
    except Exception as e:
        return {}

def extract_documents_text(directory):
    documents = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.startswith('~$'):
                continue
            path = os.path.join(root, file)
            cleaned_path = clean_path(path)
            metadata = extract_metadata(path)
            if file.endswith('.pdf'):
                text = extract_text_from_pdf(path)
                if text:
                    documents.append((cleaned_path, text, metadata))
            elif file.endswith('.docx'):
                text = extract_text_from_docx(path)
                if text:
                    documents.append((cleaned_path, text, metadata))
            elif file.endswith('.txt'):
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        text = f.read().lower()
                    documents.append((cleaned_path, text, metadata))
                except Exception as e:
                    pass
    return documents

def clean_path(path):
    return os.path.normpath(path).replace("\\", "/").lower()

def extract_training_data(training_directories):
    training_documents = []
    for label, directory in training_directories.items():
        documents = extract_documents_text(directory)
        for _, text, metadata in documents:
            training_documents.append((directory, text, label, metadata))
    return training_documents

def extract_features(documents):
    texts = [text for _, text, _ in documents]
    metadata_texts = [" ".join([str(value) for value in metadata.values()]) for _, _, metadata in documents]
    combined_texts = [f"{text} {metadata}" for text, metadata in zip(texts, metadata_texts)]
    
    vectorizer = TfidfVectorizer(stop_words=german_stopwords)
    X = vectorizer.fit_transform(combined_texts)
    return X, vectorizer

def keyword_classification(text, keyword_map):
    for label, keywords in keyword_map.items():
        for keyword in keywords:
            if keyword.lower() in text:
                return label
    return 'UNKNOWN'

def extract_entities(text):
    doc = nlp(text)
    entities = [ent.text for ent in doc.ents]
    return entities

def analyze_sentiment(text):
    results = classifier(text)
    sentiment_scores = {result['label']: result['score'] for result in results[0]}
    return sentiment_scores

def enhanced_classification(text, metadata, keyword_map):
    # Combine text and metadata
    combined_text = f"{text} {' '.join([str(value) for value in metadata.values()])}"
    
    # Keyword-basierte Klassifizierung
    label = keyword_classification(combined_text, keyword_map)
    if label != 'UNKNOWN':
        return label

    # Benannte Entitäten analysieren
    entities = extract_entities(combined_text)
    if any(entity in keyword_map.get('RED', []) for entity in entities):
        return 'RED'
    if any(entity in keyword_map.get('AMBER', []) for entity in entities):
        return 'AMBER'
    if any(entity in keyword_map.get('GREEN', []) for entity in entities):
        return 'GREEN'
    if any(entity in keyword_map.get('WHITE', []) for entity in entities):
        return 'WHITE'

    # Sentiment analysieren
    sentiment = analyze_sentiment(combined_text)
    if sentiment.get('NEGATIVE', 0) > 0.5:
        return 'RED'
    if sentiment.get('NEUTRAL', 0) > 0.5:
        return 'WHITE'
    if sentiment.get('POSITIVE', 0) > 0.5:
        return 'GREEN'

    return 'UNKNOWN'

def classify_documents_with_enhanced_techniques(documents, keyword_map):
    classified_documents = []
    unclassified_documents = []
    for path, text, metadata in documents:
        label = enhanced_classification(text, metadata, keyword_map)
        if label == 'UNKNOWN':
            unclassified_documents.append((path, text, metadata))
        else:
            classified_documents.append((path, text, metadata, label))
    return classified_documents, unclassified_documents

def train_sensitivity_model(training_documents):
    df = pd.DataFrame(training_documents, columns=['path', 'text', 'label', 'metadata'])
    # Ensure metadata is parsed correctly
    def parse_metadata(x):
        try:
            return json.loads(x) if isinstance(x, str) and x else x
        except json.JSONDecodeError:
            return {}
    
    df['metadata'] = df['metadata'].apply(parse_metadata)

    # Ensure labels are valid categories
    valid_labels = set(TLP_COLORS.keys()) - {'UNKNOWN'}
    df = df[df['label'].isin(valid_labels)]
    
    # Recreate X after filtering labels
    X = df.apply(lambda row: f"{row['text']} {' '.join([str(value) for value in row['metadata'].values()])}", axis=1)
    y = df['label']

    # Check if we have any data left after filtering
    if X.empty or y.empty:
        print("No valid training data available after filtering. Exiting.")
        return None

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    pipeline = Pipeline([
        ('tfidf', TfidfVectorizer(stop_words=german_stopwords)),
        ('clf', LogisticRegression(max_iter=200))
    ])

    pipeline.fit(X_train, y_train)
    y_pred = pipeline.predict(X_test)

    accuracy = accuracy_score(y_test, y_pred)
    print(f"Model accuracy: {accuracy}")

    return pipeline

def classify_documents(model, documents):
    combined_texts = [f"{text} {' '.join([str(value) for value in metadata.values()])}" for _, text, metadata in documents]
    labels = model.predict(combined_texts)
    return labels

def plot_clusters(X, labels, output_file=output_image):
    pca = PCA(n_components=2)
    scatter_plot_points = pca.fit_transform(X.toarray())
    x_axis = scatter_plot_points[:, 0]
    y_axis = scatter_plot_points[:, 1]
    fig, ax = plt.subplots(figsize=(20, 10))
    unique_labels = set(labels)
    for label in unique_labels:
        indices = [i for i, l in enumerate(labels) if l == label]
        ax.scatter(x_axis[indices], y_axis[indices], label=label, color=TLP_COLORS.get(label, 'black'))
    ax.legend()
    plt.savefig(output_file)

def save_classification_results(documents, labels, output_file):
    data = [(path, label) for (path, _, _), label in zip(documents, labels)]
    df = pd.DataFrame(data, columns=['Document Path', 'Sensitivity Label'])
    df.to_csv(output_file, index=False)

def main_full():
    documents = extract_documents_text(documents_dir)
    if not documents:
        return
    
    # Erweiterte Klassifizierung
    classified_documents, unclassified_documents = classify_documents_with_enhanced_techniques(documents, keyword_map)
    
    # Extract features and perform unsupervised clustering on unclassified documents
    if unclassified_documents:
        X_unclassified, vectorizer = extract_features(unclassified_documents)
        kmeans = KMeans(n_clusters=4, random_state=42).fit(X_unclassified)
        unsupervised_labels = kmeans.labels_
        
        # Add clustering results to classified documents
        for (path, text, metadata), label in zip(unclassified_documents, unsupervised_labels):
            classified_documents.append((path, text, metadata, f'CLUSTER_{label}'))
    
    # Combine training data with classified documents for model training
    training_documents = extract_training_data(training_directories)
    training_documents.extend(classified_documents)
    
    sensitivity_model = train_sensitivity_model(training_documents)
    if sensitivity_model is None:
        return

    labels = classify_documents(sensitivity_model, documents)

    save_classification_results(documents, labels, output_file)
    X, vectorizer = extract_features(documents)
    plot_clusters(X, labels)
    
    # Trigger visualize.py script
    script_path = 'visualize.py'
    try:
        subprocess.run(['python', script_path], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error running visualize.py script: {e}")

if __name__ == "__main__":
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    main_full()
